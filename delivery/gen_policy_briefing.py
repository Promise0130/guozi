# -*- coding: utf-8 -*-
"""Generate the Module C policy briefing Word document and CSV dataset."""
import csv, os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.dirname(os.path.abspath(__file__))

P = [
  dict(id="POL-001",
    title="企业国有资产交易操作规则",
    issuer="国务院国资委（产权局）", doc_date="2025-03-03",
    doc_type="正式规范性文件", doc_type_code="A",
    doc_number="修订印发（原资发产权[2009]120号）",
    topic_primary="风险防范", topic_secondary="产权管理",
    level="中央", province="",
    summary=(
      "修订完善企业国有资产交易操作规范，新增国有企业增资及实物资产转让流程，"
      "优化交易周期、降低交易成本，明确关联人回避/名称字号使用/保密义务/档案留存等要求。"
    ),
    binding_phrases="应当;不得;明确;细化;缩短",
    evidence_url="http://www.sasac.gov.cn/n2588035/n2588320/n2588335/c32955553/content.html",
    evidence_strength="S1-正式令文全文公开"),
  dict(id="POL-002",
    title="中央企业发展规划管理办法",
    issuer="国务院国资委（规划局）", doc_date="2025-06-06",
    doc_type="正式规范性文件", doc_type_code="A",
    doc_number="国务院国资委令第45号",
    topic_primary="主责主业管理", topic_secondary="内部控制",
    level="中央", province="",
    summary=(
      "首次以制度形式建立三级规划体系（总体发展规划+重点任务规划+企业规划），"
      "强调聚焦主责主业、产业优化调整，规划执行纳入考核分配与监督追责，"
      "董事长为第一责任人。6章39条。"
    ),
    binding_phrases="首次;应当;第一责任人;偏离…追究;刚性约束",
    evidence_url="http://www.sasac.gov.cn/n2588035/n2588320/n2588335/c33689673/content.html",
    evidence_strength="S1-正式令文全文公开"),
  dict(id="POL-003",
    title="关于新时代中央企业高质量推进品牌建设的意见",
    issuer="国务院国资委（社会责任局）", doc_date="2025-07-04",
    doc_type="通知/方案", doc_type_code="B",
    doc_number="",
    topic_primary="主责主业管理", topic_secondary="",
    level="中央", province="",
    summary=(
      "指导中央企业推进品牌战略与发展战略一体部署，强调六个全面任务"
      "（品牌战略/目标/过程/资产/国际化/组织保障管理），"
      "将品牌建设纳入监督考核。"
    ),
    binding_phrases="应当;完善;监督考核",
    evidence_url="http://www.sasac.gov.cn/n2588035/n2588320/n2588340/c33899498/content.html",
    evidence_strength="S1-正式文件解读全文公开"),
  dict(id="POL-004",
    title="中央企业应急管理办法（征求意见稿）",
    issuer="国务院国资委", doc_date="2025-11-12",
    doc_type="通知/方案", doc_type_code="B",
    doc_number="征求意见稿",
    topic_primary="风险防范", topic_secondary="内部控制",
    level="中央", province="",
    summary="面向社会公开征求意见，拟规范中央企业应急管理体系建设，属起草阶段。",
    binding_phrases="征求意见;拟",
    evidence_url="http://www.sasac.gov.cn/n2588035/n2588320/n2588335/c34873055/content.html",
    evidence_strength="S2-征求意见稿（尚未生效）"),
  dict(id="POL-005",
    title="中央企业违规经营投资责任追究实施办法",
    issuer="国务院国资委（监督追责局）", doc_date="2025-12-18",
    doc_type="正式规范性文件", doc_type_code="A",
    doc_number="国务院国资委令第46号",
    topic_primary="风险防范", topic_secondary="内部控制;金融业务管理;主责主业管理",
    level="中央", province="",
    summary=(
      "替代试行版（37号令），8章91条，13方面98种追责情形。"
      "新增：主责主业偏离(7.3)、内控体系缺陷(7.4)、多层架构规避监管(7.5)、"
      "融资性贸易(7.8)、金融业务6种情形(11条)、科技创新5种情形(12条)。"
      "资产损失分三档：一般(<500万)、较大(500万-5000万)、重大(>=5000万)。"
      "明确容错免责6种情形(36条)与从重处理7种情形(44条)。"
      "自2026年1月1日施行。"
    ),
    binding_phrases="应当;不得;未履行;自…施行;终身问责;禁入限制",
    evidence_url="http://www.sasac.gov.cn/n2588035/n2588320/n2588335/c35128064/content.html",
    evidence_strength="S1-正式令文全文公开（91条完整文本）"),
  dict(id="POL-006",
    title="国有企业领导人员廉洁从业规定（新修订）",
    issuer="中共中央办公厅、国务院办公厅", doc_date="2026-03-23",
    doc_type="正式规范性文件", doc_type_code="A",
    doc_number="中央两办印发",
    topic_primary="风险防范", topic_secondary="内部控制",
    level="中央", province="",
    summary=(
      "全面修订2009年版，适用范围扩大至国有全资企业和实际控制企业，"
      "聚焦关键少数（3类领导人员），"
      "强调对并购重组/产权交易/招标投标等关键环节加大监督，"
      "严查靠企吃企/设租寻租/利益输送。"
      "完善经济责任信息报告制度。"
    ),
    binding_phrases="不得;应当;严肃处理;依法追究刑事责任;解任;解聘",
    evidence_url="http://www.sasac.gov.cn/n2588025/n2588119/c35382734/content.html",
    evidence_strength="S1-答记者问全文公开（原文由中央两办印发）"),
  dict(id="POL-007",
    title="关于推动中央企业加快财务数智化转型升级的指导意见",
    issuer="国务院国资委（财管运行局）", doc_date="2026-03-25",
    doc_type="通知/方案", doc_type_code="B",
    doc_number="",
    topic_primary="司库体系建设", topic_secondary="内部控制",
    level="中央", province="",
    summary=(
      "五大升级方向：财务管理体系化、财务系统集成化、财务信息数字化、"
      "财务监督智能化、财务转型一体化。"
      "目标建成全经营领域的数字化资源管理平台（DRP系统），"
      "充分应用大数据/大模型/人工智能等新技术推动穿透监管。"
    ),
    binding_phrases="应用大数据/大模型/AI;五大升级方向;穿透监管;DRP",
    evidence_url="http://www.sasac.gov.cn/n2588025/n2643314/c35387473/content.html",
    evidence_strength="S3-新闻通稿引用文件名（全文未公开）"),
  dict(id="POL-008",
    title="黑龙江省国资委2025年度制度建设成果（35项）",
    issuer="黑龙江省国资委", doc_date="2025-12-31",
    doc_type="通知/方案", doc_type_code="B",
    doc_number="省级系列文件",
    topic_primary="内部控制", topic_secondary="风险防范",
    level="省级", province="黑龙江",
    summary=(
      "2025年累计出台35项制度文件，含：省国资委监管企业法律纠纷案件管理实施办法、"
      "关于加强出资企业内部审计监督工作的实施意见、"
      "省国资委监管企业违规经营投资责任追究实施办法、"
      "省国资委出资企业审计问题整改标准指引(1.0版)。"
      "首次集中宣讲会（360余人）。建立立改废动态调整机制。"
    ),
    binding_phrases="实施办法;实施意见;指引;立改废;闭环管控",
    evidence_url="http://www.sasac.gov.cn/n2588025/n2588129/c35387101/content.html",
    evidence_strength="S3-新闻通稿（省级制度全文未公开）"),
  dict(id="POL-009",
    title="国务院国资委召开中央企业法治工作会议",
    issuer="国务院国资委", doc_date="2026-04-01",
    doc_type="工作要点", doc_type_code="C",
    doc_number="",
    topic_primary="内部控制", topic_secondary="风险防范",
    level="中央", province="",
    summary="2026年4月1日召开中央企业法治工作会议，部署法治建设与合规管理工作。",
    binding_phrases="部署;会议",
    evidence_url="http://www.sasac.gov.cn/n2588025/n2643314/c35392672/content.html",
    evidence_strength="S4-标题可见但正文未获取"),
  dict(id="POL-010",
    title="浙江省2026年国资国企党风廉政建设和反腐败工作会议暨中央巡视反馈国企问题整改推进会",
    issuer="浙江省国资委", doc_date="2026-03-16",
    doc_type="工作要点", doc_type_code="C",
    doc_number="",
    topic_primary="风险防范", topic_secondary="内部控制",
    level="省级", province="浙江",
    summary="副省长张雁云出席并讲话，部署2026年党风廉政建设与中央巡视整改。",
    binding_phrases="部署;整改;讲话",
    evidence_url="https://gzw.zj.gov.cn/col/col1229430728/art/2026/art_95d768027391415bbb62b4d06b889d4d.html",
    evidence_strength="S3-新闻通稿（省级，详情页未获取）"),
  dict(id="POL-011",
    title="福建省国资委实施五大行动推动所出资企业向新向优发展",
    issuer="福建省国资委", doc_date="2026-03-27",
    doc_type="工作要点", doc_type_code="C",
    doc_number="",
    topic_primary="主责主业管理", topic_secondary="",
    level="省级", province="福建",
    summary="福建省国资委部署五大行动推动向新向优发展，具体内容涉及主责主业聚焦等。",
    binding_phrases="行动;推动",
    evidence_url="http://www.sasac.gov.cn/n2588025/n2588129/c35387981/content.html",
    evidence_strength="S3-新闻通稿"),
  dict(id="POL-012",
    title="宁波市国资委部署三提行动推动国资国企高质量发展",
    issuer="宁波市国资委", doc_date="2026-04-01",
    doc_type="工作要点", doc_type_code="C",
    doc_number="",
    topic_primary="主责主业管理", topic_secondary="",
    level="地级市", province="浙江",
    summary="宁波市国资委部署三提行动（提质/提速/提效），推动国资国企高质量发展。",
    binding_phrases="行动;部署",
    evidence_url="http://www.sasac.gov.cn/n2588025/n2588129/c35391405/content.html",
    evidence_strength="S3-新闻通稿"),
]

# ── CSV ─────────────────────────────────────────────────────────────
def write_csv():
    path = os.path.join(OUT, "guozi_policy_v1_20260402.csv")
    fields = list(P[0].keys())
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        for p in P:
            w.writerow(p)
    print(f"CSV: {path}  ({len(P)} rows)")

# ── DOCX ────────────────────────────────────────────────────────────
def hd(doc, txt, lv):
    h = doc.add_heading(txt, level=lv)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

def ap(doc, txt, bold=False, sz=11):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.bold = bold
    return p

def pol_sec(doc, pol):
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{pol['id']}] ")
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(100, 100, 100)
    r2 = p.add_run(pol["title"])
    r2.bold = True
    r2.font.size = Pt(12)
    # meta table
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta = [
        ("发文机关", pol["issuer"]),
        ("发文日期", pol["doc_date"]),
        ("文种分级", f"{pol['doc_type']}（{pol['doc_type_code']}级）"),
        ("文号", pol["doc_number"] or "—"),
        ("主题", pol["topic_primary"] + (f" / {pol['topic_secondary']}" if pol["topic_secondary"] else "")),
    ]
    for i, (k, v) in enumerate(meta):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
        for c in tbl.rows[i].cells:
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)
        tbl.rows[i].cells[0].width = Cm(3)
        tbl.rows[i].cells[1].width = Cm(13)
    doc.add_paragraph()
    ap(doc, "核心内容：", bold=True)
    ap(doc, pol["summary"])
    ap(doc, f"约束性表述关键词：{pol['binding_phrases']}")
    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"证据强度：{pol['evidence_strength']}")
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(100, 100, 100)
    p4 = doc.add_paragraph()
    r4 = p4.add_run(f"原文链接：{pol['evidence_url']}")
    r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor(0, 0, 180)
    doc.add_paragraph()

def write_docx():
    doc = Document()
    t = doc.add_heading("2025年以来国资管理政策动态简报", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0, 51, 102)
    ap(doc, f"编制日期：{date.today().isoformat()}    编制方：自动采集+人工校核")
    ap(doc, "聚焦领域：主责主业管理 | 司库体系建设 | 内部控制 | 金融业务管理 | 风险防范")
    ap(doc, f"收录政策条目：{len(P)} 条", bold=True)
    doc.add_page_break()
    # TOC
    hd(doc, "目 录", 1)
    for pol in P:
        ap(doc, f"{pol['id']}  {pol['title']}  [{pol['doc_type']}]  {pol['doc_date']}")
    doc.add_page_break()
    # Section A
    hd(doc, "一、正式规范性文件", 1)
    ap(doc, "以部委令、两办印发等形式发布的正式制度文件，对被监管企业具有直接约束力。")
    doc.add_paragraph()
    for pol in P:
        if pol["doc_type_code"] == "A":
            pol_sec(doc, pol)
    doc.add_page_break()
    # Section B
    hd(doc, "二、通知 / 方案 / 指导意见", 1)
    ap(doc, "以通知、指导意见、征求意见稿等形式发布，具有方向性约束但执行弹性较大。")
    doc.add_paragraph()
    for pol in P:
        if pol["doc_type_code"] == "B":
            pol_sec(doc, pol)
    doc.add_page_break()
    # Section C
    hd(doc, "三、工作要点 / 会议部署", 1)
    ap(doc, "年度工作要点、专题会议等政策信号，为工作部署性质。注意区分与一般宣传报道的差异。")
    doc.add_paragraph()
    for pol in P:
        if pol["doc_type_code"] == "C":
            pol_sec(doc, pol)
    doc.add_page_break()
    # Section: evidence levels
    hd(doc, "四、证据强度分级说明", 1)
    lvs = [
        ("S1","正式令文/规定全文公开","可直接引用条款原文"),
        ("S2","征求意见稿/草案公开","尚未生效，引用时须标注拟"),
        ("S3","新闻通稿/答记者问","可引用文件名和关键表述，无法逐条核实"),
        ("S4","标题可见但正文未获取","仅可引用标题和日期，内容待补"),
    ]
    tbl = doc.add_table(rows=len(lvs)+1, cols=3)
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(["等级","定义","引用建议"]):
        tbl.rows[0].cells[j].text = h
    for i, (c, d, a) in enumerate(lvs):
        tbl.rows[i+1].cells[0].text = c
        tbl.rows[i+1].cells[1].text = d
        tbl.rows[i+1].cells[2].text = a
    doc.add_paragraph()
    # Section: doc type classification
    hd(doc, "五、文种分级说明", 1)
    dts = [
        ("A","正式规范性文件","部委令、两办印发、条例、办法","直接约束，必须执行"),
        ("B","通知/方案/指导意见","关于…的通知、指导意见、征求意见稿","方向性约束"),
        ("C","工作要点/会议部署","年度工作会议、专题推进会","信号性，无直接条款约束力"),
        ("D","领导讲话","讲话稿、署名文章","政治信号，通常无条款"),
        ("E","新闻通稿/宣传报道","地方扫描、媒体报道","最弱，仅供参考"),
    ]
    tbl2 = doc.add_table(rows=len(dts)+1, cols=4)
    tbl2.style = "Light Grid Accent 1"
    for j, h in enumerate(["级别","文种名称","典型形式","约束力判断"]):
        tbl2.rows[0].cells[j].text = h
    for i, (c, n, f, b) in enumerate(dts):
        tbl2.rows[i+1].cells[0].text = c
        tbl2.rows[i+1].cells[1].text = n
        tbl2.rows[i+1].cells[2].text = f
        tbl2.rows[i+1].cells[3].text = b
    doc.add_paragraph()
    # Section: topic stats
    hd(doc, "六、主题覆盖统计", 1)
    topics = {}
    for pol in P:
        for tt in [pol["topic_primary"]] + pol["topic_secondary"].split(";"):
            tt = tt.strip()
            if tt:
                topics[tt] = topics.get(tt, 0) + 1
    for tt, cnt in sorted(topics.items(), key=lambda x: -x[1]):
        ap(doc, f"  {tt}：{cnt} 条")
    doc.add_paragraph()
    ap(doc, "【简报结束】", bold=True)
    path = os.path.join(OUT, "guozi_policy_briefing_v1_20260402.docx")
    doc.save(path)
    print(f"DOCX: {path}")

if __name__ == "__main__":
    write_csv()
    write_docx()
    print("Done.")
